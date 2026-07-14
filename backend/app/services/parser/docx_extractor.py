"""DOCX text extraction using python-docx."""
import io
import re
import docx
from app.core.exceptions import ValidationError


def extract_docx_text(file_bytes: bytes) -> dict:
    """Extract plain text, lists, and tables from DOCX file bytes.

    Returns:
        dict: {
            "text": str,
            "page_count": int,
            "character_count": int,
            "word_count": int
        }
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValidationError(
            "Could not open DOCX file. It may be corrupt or encrypted.",
            details=f"CORRUPT_DOCX: {str(e)}"
        )

    lines = []

    # 1. Process paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Check if this is a list item and make sure it has a bullet prefix
        style_name = p.style.name.lower() if p.style and p.style.name else ""
        if any(kw in style_name for kw in ["list", "bullet", "desc"]):
            # Prepend bullet if not already containing list item markers
            if not text.startswith(('•', '-', '*', 'o', '1.', '2.', '3.')):
                text = "• " + text

        lines.append(text)

    # 2. Process tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            last_cell_text = None
            for cell in row.cells:
                cell_txt = cell.text.strip()
                # Deduplicate merged cells text
                if cell_txt and cell_txt != last_cell_text:
                    row_text.append(cell_txt)
                    last_cell_text = cell_txt
            if row_text:
                lines.append(" | ".join(row_text))

    # Join and normalize whitespace
    combined = "\n".join(lines)
    normalized_lines = []
    for line in combined.splitlines():
        cleaned_line = line.strip()
        cleaned_line = re.sub(r' +', ' ', cleaned_line)
        if cleaned_line:
            normalized_lines.append(cleaned_line)

    normalized_text = "\n".join(normalized_lines)
    character_count = len(normalized_text)
    word_count = len(normalized_text.split())

    # DOCX does not natively store page count without rendering.
    # We can estimate based on character/word density, or default to 1.
    # A standard page has ~500 words or ~3000 characters.
    estimated_pages = max(1, (word_count // 450) + 1)

    return {
        "text": normalized_text,
        "page_count": estimated_pages,
        "character_count": character_count,
        "word_count": word_count,
    }
