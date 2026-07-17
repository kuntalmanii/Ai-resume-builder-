"""Integration tests for Resume Import Sessions, secure validation, parsing, and Career Profile imports."""
import io
import uuid
from datetime import UTC, datetime, timedelta

import docx
import fitz
import pytest
from httpx import AsyncClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models.career_entry import CareerEntry
from app.db.models.resume_import_session import ResumeImportSession
from app.db.models.resume_version import ResumeVersion

pytestmark = pytest.mark.asyncio


# ─── Helpers ──────────────────────────────────────────────────────────────────
async def _register_and_login(client: AsyncClient, email: str, name: str) -> str:
    await client.post(
        "/api/v1/auth/register",
        json={"full_name": name, "email": email, "password": "Testpass1!"},
    )
    res = await client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": "Testpass1!"},
    )
    return res.json()["access_token"]


def _create_test_pdf_bytes(text_content: str) -> bytes:
    """Generate a valid PDF in memory using PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text_content)
    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes


def _create_test_docx_bytes(paragraphs: list[tuple[str, str | None]]) -> bytes:
    """Generate a valid DOCX in memory using python-docx."""
    doc = docx.Document()
    for text, style in paragraphs:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    docx_io = io.BytesIO()
    doc.save(docx_io)
    return docx_io.getvalue()


# ─── Tests ────────────────────────────────────────────────────────────────────

async def test_unauthenticated_upload_rejected(client: AsyncClient) -> None:
    """Upload without bearer token is rejected with 401."""
    pdf_bytes = _create_test_pdf_bytes("John Doe Resume")
    files = {"file": ("resume.pdf", pdf_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", files=files)
    assert res.status_code == 401


async def test_valid_pdf_upload_accepted(client: AsyncClient) -> None:
    """Upload a valid PDF resume, verify parsing and session creation."""
    token = await _register_and_login(client, "test1@example.com", "Test User")
    headers = {"Authorization": f"Bearer {token}"}

    pdf_text = (
        "John Doe\n"
        "Email: john.doe@example.com | Phone: +1-555-555-0199\n"
        "Website: https://johndoe.dev | GitHub: https://github.com/johndoe\n"
        "Professional Summary\n"
        "Experienced software developer.\n"
        "Education\n"
        "University of Tech - Bachelor of Science in CS (2018 - 2022)\n"
        "Experience\n"
        "Acme Corp - Senior Engineer (2022 - Present)\n"
        "• Built awesome cloud platform.\n"
        "Skills\n"
        "Python, FastAPI, Postgres, Docker\n"
    )
    pdf_bytes = _create_test_pdf_bytes(pdf_text)
    files = {"file": ("my_resume.pdf", pdf_bytes, "application/pdf")}

    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 201
    data = res.json()

    assert data["id"] is not None
    assert data["original_filename"] == "my_resume.pdf"
    assert data["document_type"] == "pdf"
    assert data["status"] == "review_ready"

    # Verify parsing output
    doc = data["parsed_document"]
    assert doc["personal_information"]["full_name"] == "John Doe"
    assert doc["personal_information"]["email"] == "john.doe@example.com"
    assert doc["personal_information"]["phone"] == "+1-555-555-0199"
    assert doc["personal_information"]["github_url"] == "https://github.com/johndoe"
    assert doc["professional_summary"] == "Experienced software developer."

    # Check sections detected
    assert "education" in data["detected_sections"]
    assert "experience" in data["detected_sections"]
    assert "skills" in data["detected_sections"]
    assert "certifications" in data["missing_sections"]


async def test_valid_docx_upload_accepted(client: AsyncClient) -> None:
    """Upload a valid DOCX resume, verify parsing and session creation."""
    token = await _register_and_login(client, "test2@example.com", "Test User")
    headers = {"Authorization": f"Bearer {token}"}

    paras = [
        ("Jane Smith", None),
        ("Email: jane@example.com | Phone: 9876543210", None),
        ("Education", "Heading 1"),
        ("MIT - Bachelor of Science in Physics", None),
        ("Experience", "Heading 1"),
        ("SpaceX - Propulsion Engineer", None),
        ("Worked on rocket landing software.", "List Bullet"),
    ]
    docx_bytes = _create_test_docx_bytes(paras)
    files = {"file": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 201
    data = res.json()
    assert data["original_filename"] == "resume.docx"
    assert data["document_type"] == "docx"
    doc = data["parsed_document"]
    assert doc["personal_information"]["email"] == "jane@example.com"


async def test_unsupported_extensions_and_fake_signatures_rejected(client: AsyncClient) -> None:
    """Verifies that extension mismatch, fake extensions, and corrupt files are rejected."""
    token = await _register_and_login(client, "test3@example.com", "Test User")
    headers = {"Authorization": f"Bearer {token}"}

    # 1. Unsupported extension
    files = {"file": ("resume.txt", b"some plain text", "text/plain")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 400
    assert res.json()["error"]["code"] == "VALIDATION_ERROR"

    # 2. Fake PDF (PDF extension but random bytes)
    files = {"file": ("resume.pdf", b"fake pdf bytes that do not start with pdf", "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 400
    assert "INVALID_PDF_SIGNATURE" in res.json()["error"]["details"]

    # 3. Fake DOCX (DOCX extension but random bytes)
    files = {"file": ("resume.docx", b"fake zip bytes", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 400
    assert "INVALID_DOCX_SIGNATURE" in res.json()["error"]["details"]

    # 4. Empty file
    files = {"file": ("resume.pdf", b"", "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 400
    assert "EMPTY_FILE" in res.json()["error"]["details"]


async def test_oversized_file_rejected(client: AsyncClient) -> None:
    """Verifies that files exceeding size limit are rejected with FILE_TOO_LARGE."""
    token = await _register_and_login(client, "test4@example.com", "Test User")
    headers = {"Authorization": f"Bearer {token}"}

    # Exceeds 10MB by creating huge bytes array
    large_bytes = b"%PDF" + (b"0" * (11 * 1024 * 1024))
    files = {"file": ("large.pdf", large_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)

    assert res.status_code == 400
    err = res.json()["error"]
    assert err["code"] == "FILE_TOO_LARGE"
    assert err["details"]["max_size_mb"] == 10


async def test_scanned_pdf_ocr_required(client: AsyncClient) -> None:
    """Verifies that PDF with pages but zero text triggers OCR required warning."""
    token = await _register_and_login(client, "test5@example.com", "Test User")
    headers = {"Authorization": f"Bearer {token}"}

    # Create PDF page with no text (empty canvas)
    doc = fitz.open()
    doc.new_page() # no page text
    pdf_bytes = doc.write()
    doc.close()

    files = {"file": ("scanned.pdf", pdf_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    assert res.status_code == 400
    assert "SCANNED_PDF_OCR_REQUIRED" in res.json()["error"]["details"]


async def test_ownership_enforcement(client: AsyncClient, db_session: AsyncSession) -> None:
    """Verify that users can only get, update, and finalize their own sessions."""
    token1 = await _register_and_login(client, "alice@example.com", "Alice")
    token2 = await _register_and_login(client, "bob@example.com", "Bob")

    headers1 = {"Authorization": f"Bearer {token1}"}
    headers2 = {"Authorization": f"Bearer {token2}"}

    # Alice uploads
    pdf_bytes = _create_test_pdf_bytes("Alice Resume text. " * 10)
    files = {"file": ("alice_resume.pdf", pdf_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers1, files=files)
    import_id = res.json()["id"]

    # Bob tries to access Alice's session -> 404
    res_get = await client.get(f"/api/v1/resume-imports/{import_id}", headers=headers2)
    assert res_get.status_code == 404

    # Bob tries to edit Alice's session -> 404
    res_patch = await client.patch(
        f"/api/v1/resume-imports/{import_id}/document",
        headers=headers2,
        json={"parsed_document": res.json()["parsed_document"]}
    )
    assert res_patch.status_code == 404

    # Bob tries to finalize Alice's session -> 404
    res_final = await client.post(
        f"/api/v1/resume-imports/{import_id}/finalize",
        headers=headers2,
        json={"title": "Bob's Resume", "template_id": "modern", "import_to_career_profile": False}
    )
    assert res_final.status_code == 404


async def test_expired_session_rejected(client: AsyncClient, db_session: AsyncSession) -> None:
    """Expired session update or finalization is rejected."""
    token = await _register_and_login(client, "expired@example.com", "Expired User")
    headers = {"Authorization": f"Bearer {token}"}

    # Upload
    pdf_bytes = _create_test_pdf_bytes("Expired Session Resume. " * 10)
    files = {"file": ("res.pdf", pdf_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    import_id = res.json()["id"]

    # Manually expire the session in the database
    stmt = select(ResumeImportSession).where(ResumeImportSession.id == uuid.UUID(import_id))
    db_res = await db_session.execute(stmt)
    session = db_res.scalar_one()
    session.expires_at = datetime.now(UTC) - timedelta(minutes=10)
    db_session.add(session)
    await db_session.commit()

    # Try to finalize -> should fail
    res_final = await client.post(
        f"/api/v1/resume-imports/{import_id}/finalize",
        headers=headers,
        json={"title": "Expired Resume", "template_id": "modern", "import_to_career_profile": False}
    )
    assert res_final.status_code == 400
    assert "SESSION_EXPIRED" in res_final.json()["error"]["details"]


async def test_update_parsed_document_validation(client: AsyncClient) -> None:
    """Can update parsed document with valid JSON schema, but invalid schemas are rejected."""
    token = await _register_and_login(client, "validator@example.com", "Val User")
    headers = {"Authorization": f"Bearer {token}"}

    pdf_bytes = _create_test_pdf_bytes("Testing validation. " * 10)
    files = {"file": ("res.pdf", pdf_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    import_id = res.json()["id"]
    doc = res.json()["parsed_document"]

    # Update with valid change
    doc["personal_information"]["full_name"] = "New Valid Name"
    res_patch = await client.patch(
        f"/api/v1/resume-imports/{import_id}/document",
        headers=headers,
        json={"parsed_document": doc}
    )
    assert res_patch.status_code == 200
    assert res_patch.json()["parsed_document"]["personal_information"]["full_name"] == "New Valid Name"

    # Update with invalid structure (missing education structure, etc.)
    invalid_doc = doc.copy()
    invalid_doc["education"] = "not a list"
    res_patch_invalid = await client.patch(
        f"/api/v1/resume-imports/{import_id}/document",
        headers=headers,
        json={"parsed_document": invalid_doc}
    )
    assert res_patch_invalid.status_code == 422
    assert res_patch_invalid.json()["error"]["code"] == "VALIDATION_ERROR"


async def test_finalize_creates_resume_version_and_profile_entries(client: AsyncClient, db_session: AsyncSession) -> None:
    """Finalization creates Resume, ResumeVersion snapshot, CareerProfile entries, and prevents double finalization."""
    token = await _register_and_login(client, "finalizer@example.com", "Finalizer User")
    headers = {"Authorization": f"Bearer {token}"}

    # Upload
    pdf_text = (
        "Alice Smith\n"
        "Email: alice@example.com\n"
        "Education\n"
        "Harvard - Bachelor of Arts (2015 - 2019)\n"
        "Experience\n"
        "Google - Software Engineer (2019 - Present)\n"
        "• Developed features for search.\n"
        "Skills\n"
        "Go, C++, Java\n"
    )
    pdf_bytes = _create_test_pdf_bytes(pdf_text)
    files = {"file": ("alice.pdf", pdf_bytes, "application/pdf")}
    res = await client.post("/api/v1/resume-imports", headers=headers, files=files)
    import_id = res.json()["id"]

    # Finalize with Career Profile Import
    res_final = await client.post(
        f"/api/v1/resume-imports/{import_id}/finalize",
        headers=headers,
        json={
            "title": "My Awesome Resume",
            "template_id": "creative",
            "import_to_career_profile": True,
            "selected_entries": ["education", "experience", "skills"]
        }
    )
    assert res_final.status_code == 200
    resume_data = res_final.json()

    assert resume_data["title"] == "My Awesome Resume"
    assert resume_data["template_id"] == "creative"
    assert resume_data["source_type"] == "uploaded_pdf"
    assert resume_data["original_filename"] == "alice.pdf"
    # Ensure no internal filesystem paths are exposed in response
    assert "original_file_path" not in resume_data or resume_data["original_file_path"] is None

    resume_uuid = uuid.UUID(resume_data["id"])

    # 1. Verify Resume version snapshot is created
    v_stmt = select(ResumeVersion).where(ResumeVersion.resume_id == resume_uuid)
    v_res = await db_session.execute(v_stmt)
    versions = v_res.scalars().all()
    assert len(versions) == 1
    assert versions[0].version_number == 1
    assert versions[0].change_reason == "Initial import snapshot"

    # 2. Verify Career Profile Entries are created with correct verification status
    c_stmt = select(CareerEntry).where(CareerEntry.source_type == "resume_import")
    c_res = await db_session.execute(c_stmt)
    entries = c_res.scalars().all()

    # Harvard (edu) + Google (experience) + Go, C++, Java (3 skills) = 5 entries
    assert len(entries) == 5
    for entry in entries:
        assert entry.verification_status == "user_confirmed"

    # Verify duplicate prevention: call finalize again (should be rejected since session is finalized)
    res_double = await client.post(
        f"/api/v1/resume-imports/{import_id}/finalize",
        headers=headers,
        json={"title": "Second Time", "template_id": "modern", "import_to_career_profile": False}
    )
    assert res_double.status_code == 400
    assert "ALREADY_FINALIZED" in res_double.json()["error"]["details"]
